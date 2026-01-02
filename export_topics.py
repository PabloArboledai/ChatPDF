#!/usr/bin/env python3
"""Exportador multi-formato por temas desde un PDF.

Formatos soportados:
- md: Markdown legible con imágenes (no pixel-perfect).
- html: HTML con CSS inline + imágenes.
- docx: Word editable con imágenes.
- pdf: PDF por tema (pixel-perfect) copiando páginas del PDF original.
- txt: Texto plano.
- json: Estructura serializada (útil para pipelines/IA).

Notas de fidelidad:
- Si tu prioridad es “idéntico visual”, usa `pdf` (divide por rangos de páginas).
- Si tu prioridad es “editable + legible”, usa `docx` / `md` / `html`.
"""

from __future__ import annotations

import argparse
import json
import os
import re
import statistics
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF

try:
    from docx import Document
    from docx.shared import Inches
except Exception:
    Document = None  # type: ignore
    Inches = None  # type: ignore

try:
    from docx.image.exceptions import UnrecognizedImageError
except Exception:
    UnrecognizedImageError = Exception  # type: ignore

try:
    from PIL import Image
except Exception:
    Image = None  # type: ignore


_BULLET_RE = re.compile(r"^\s*([•◦·\-–])\s+(.*)$")
_ENUM_RE = re.compile(r"^\s*(\d{1,3})[\.)]\s+(.*)$")


def slugify(value: str, max_len: int = 80) -> str:
    value = value.strip()
    value = re.sub(r"\s+", " ", value)
    value = value.lower()
    value = re.sub(r"[^a-z0-9áéíóúñü\s\-]", "", value)
    value = value.replace(" ", "-")
    value = re.sub(r"-+", "-", value).strip("-")
    return value[:max_len] or "tema"


def normalize_title(s: str) -> str:
    return re.sub(r"\s+", " ", s.replace("\t", " ")).strip()


def is_probably_footer_or_noise(line: str) -> bool:
    s = line.strip()
    if not s:
        return True
    if len(s) <= 2 and s.isdigit():
        return True
    if "creative commons" in s.lower() or "licencia" in s.lower():
        return True
    letters = sum(ch.isalpha() for ch in s)
    if letters > 0 and letters / max(1, len(s)) < 0.25:
        return True
    return False


def looks_like_page_number(line: str) -> bool:
    s = line.strip()
    return bool(re.fullmatch(r"\d{1,4}", s))


@dataclass
class Item:
    kind: str  # 'text' | 'image' | 'callout'
    y0: float
    x0: float
    y1: float
    x1: float
    text: Optional[str] = None
    font_size: Optional[float] = None
    image_bytes: Optional[bytes] = None
    image_ext: Optional[str] = None


def _median(values: list[float]) -> float:
    if not values:
        return 0.0
    return float(statistics.median(values))


def collect_items(doc: fitz.Document, page_index: int) -> list[Item]:
    page = doc[page_index]
    d = page.get_text("dict")

    items: list[Item] = []
    font_sizes: list[float] = []

    for block in d.get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            spans = line.get("spans", [])
            text = "".join(s.get("text", "") for s in spans).strip()
            if not text:
                continue
            size = max((float(s.get("size", 0.0)) for s in spans), default=0.0)
            bbox = line.get("bbox") or block.get("bbox")
            if not bbox:
                continue
            font_sizes.append(size)
            items.append(
                Item(
                    kind="text",
                    x0=float(bbox[0]),
                    y0=float(bbox[1]),
                    x1=float(bbox[2]),
                    y1=float(bbox[3]),
                    text=text,
                    font_size=size,
                )
            )

    img_seen = set()
    for img in page.get_images(full=True):
        xref = img[0]
        if xref in img_seen:
            continue
        rects = page.get_image_rects(xref)
        if not rects:
            continue
        try:
            info = doc.extract_image(xref)
        except Exception:
            continue
        img_bytes = info.get("image")
        img_ext = info.get("ext", "png")
        img_seen.add(xref)
        for r in rects:
            items.append(
                Item(
                    kind="image",
                    x0=float(r.x0),
                    y0=float(r.y0),
                    x1=float(r.x1),
                    y1=float(r.y1),
                    image_bytes=img_bytes,
                    image_ext=img_ext,
                )
            )

    # “Recuadros” aproximados: rectángulos de dibujos
    rects: list[fitz.Rect] = []
    try:
        for dr in page.get_drawings():
            r = dr.get("rect")
            if r is None:
                continue
            rr = fitz.Rect(r)
            if rr.get_area() < 10_000:
                continue
            if rr.width > page.rect.width * 0.98 and rr.height > page.rect.height * 0.98:
                continue
            rects.append(rr)
    except Exception:
        rects = []

    if rects:
        consumed_ids: set[int] = set()
        for rr in rects:
            contained: list[Item] = []
            for idx, it in enumerate(items):
                if it.kind != "text" or idx in consumed_ids:
                    continue
                if it.x0 >= rr.x0 - 2 and it.x1 <= rr.x1 + 2 and it.y0 >= rr.y0 - 2 and it.y1 <= rr.y1 + 2:
                    contained.append(it)
                    consumed_ids.add(idx)
            if contained:
                contained_sorted = sorted(contained, key=lambda t: (t.y0, t.x0))
                body = "\n".join(t.text for t in contained_sorted if t.text)
                items.append(
                    Item(
                        kind="callout",
                        x0=float(rr.x0),
                        y0=float(rr.y0),
                        x1=float(rr.x1),
                        y1=float(rr.y1),
                        text=body.strip(),
                        font_size=_median(font_sizes) if font_sizes else None,
                    )
                )
        if consumed_ids:
            items = [it for idx, it in enumerate(items) if idx not in consumed_ids]

    items.sort(key=lambda it: (it.y0, it.x0))
    return items


def find_printed_page_number(items: list[Item], page_height: float) -> Optional[int]:
    candidates: list[tuple[float, int]] = []
    for it in items:
        if it.kind != "text" or not it.text:
            continue
        s = it.text.strip()
        if not looks_like_page_number(s):
            continue
        try:
            n = int(s)
        except Exception:
            continue
        if it.y0 >= page_height * 0.80:
            candidates.append((it.y0, n))
    if not candidates:
        return None
    candidates.sort(key=lambda t: t[0], reverse=True)
    return candidates[0][1]


def estimate_printed_to_pdf_offset(doc: fitz.Document, scan_pages: int = 80) -> Optional[int]:
    offsets: list[int] = []
    for page_i in range(min(scan_pages, doc.page_count)):
        page = doc[page_i]
        items = collect_items(doc, page_i)
        printed = find_printed_page_number(items, page.rect.height)
        if printed is None:
            continue
        offsets.append((page_i + 1) - printed)
    if not offsets:
        return None
    return int(statistics.median(offsets))


def parse_toc_entries(doc: fitz.Document, max_pages: int = 40) -> list[tuple[str, int]]:
    entries: list[tuple[str, int]] = []
    toc_re = re.compile(
        r"^\s*(\d{1,2})\s*[\.)]\s*(.+?)\s*[\.·…\s]{2,}(\d{1,4})\s*$",
        re.IGNORECASE,
    )
    for page_i in range(min(max_pages, doc.page_count)):
        items = collect_items(doc, page_i)
        lines = [it.text.strip() for it in items if it.kind == "text" and it.text]
        for ln in lines:
            if is_probably_footer_or_noise(ln):
                continue
            m = toc_re.match(ln)
            if not m:
                continue
            title = normalize_title(m.group(2))
            try:
                pn = int(m.group(3))
            except Exception:
                continue
            if len(title) < 3:
                continue
            entries.append((title, pn))

    seen = set()
    dedup: list[tuple[str, int]] = []
    for t, p in entries:
        key = (t.lower(), p)
        if key in seen:
            continue
        seen.add(key)
        dedup.append((t, p))
    return dedup


def format_text_line(line: str) -> tuple[str, str]:
    """Devuelve (kind, text) donde kind es 'li' o 'p'."""
    m = _BULLET_RE.match(line)
    if m:
        return "li", m.group(2).strip()
    m = _ENUM_RE.match(line)
    if m:
        return "li", f"{m.group(1)}. {m.group(2).strip()}"
    return "p", line.strip()


def heading_level(size: float, base_size: float, heading_scale: float) -> int:
    if base_size <= 0:
        return 0
    if size >= base_size * (heading_scale + 0.6):
        return 1
    if size >= base_size * heading_scale:
        return 2
    return 0


def extract_topic_assets(
    doc: fitz.Document,
    page_range: range,
    topic_dir: Path,
    heading_scale: float,
) -> tuple[list[dict], list[dict]]:
    """Extrae elementos estructurados y asegura que las imágenes estén en disco.

    Retorna (elements, images_metadata)
    elements: lista de dicts con type: heading/paragraph/list/callout/image
    """
    elements: list[dict] = []
    images_meta: list[dict] = []

    for page_i in page_range:
        page = doc[page_i]
        items = collect_items(doc, page_i)
        sizes = [it.font_size for it in items if it.kind == "text" and it.font_size]
        base = _median([float(s) for s in sizes if s]) if sizes else 0.0

        elements.append({"type": "page_marker", "page": page_i + 1})

        current_list: list[str] = []

        def flush_list():
            nonlocal current_list
            if current_list:
                elements.append({"type": "list", "items": current_list[:]})
                current_list = []

        for it in items:
            if it.kind == "image":
                flush_list()
                img_dir = topic_dir / "images"
                img_dir.mkdir(parents=True, exist_ok=True)
                img_name = f"p{page_i+1:03d}_y{int(it.y0):04d}_x{int(it.x0):04d}.{it.image_ext or 'png'}"
                img_path = img_dir / img_name
                if it.image_bytes and not img_path.exists():
                    img_path.write_bytes(it.image_bytes)
                rel = f"images/{img_name}"
                elements.append({"type": "image", "src": rel})
                images_meta.append({"page": page_i + 1, "src": rel})
                continue

            if it.kind == "callout" and it.text:
                flush_list()
                body_lines = [ln.strip() for ln in it.text.splitlines() if ln.strip()]
                body_lines = [ln for ln in body_lines if not is_probably_footer_or_noise(ln)]
                if body_lines:
                    elements.append({"type": "callout", "kind": "note", "lines": body_lines})
                continue

            if it.kind != "text" or not it.text:
                continue
            raw = it.text.strip()
            if is_probably_footer_or_noise(raw):
                continue

            lvl = heading_level(float(it.font_size or 0.0), base, heading_scale)
            if lvl:
                flush_list()
                elements.append({"type": "heading", "level": lvl + 1, "text": normalize_title(raw)})
                continue

            kind, text = format_text_line(raw)
            if kind == "li":
                current_list.append(text)
            else:
                flush_list()
                elements.append({"type": "paragraph", "text": text})

        flush_list()

    return elements, images_meta


def write_md(topic_dir: Path, title: str, elements: list[dict]) -> Path:
    topic_dir.mkdir(parents=True, exist_ok=True)
    p = topic_dir / "tema.md"
    lines: list[str] = [f"# {title}", ""]
    for el in elements:
        t = el["type"]
        if t == "page_marker":
            lines.append(f"\n<!-- Página {el['page']} -->\n")
        elif t == "heading":
            level = int(el.get("level", 2))
            level = max(2, min(6, level))
            lines.append("#" * level + " " + el["text"])
        elif t == "paragraph":
            lines.append(el["text"])
        elif t == "list":
            for it in el["items"]:
                # soporta "1. ..." ya formateado
                if re.match(r"^\d+\.\s+", it):
                    lines.append(it)
                else:
                    lines.append(f"- {it}")
        elif t == "callout":
            lines.append("> [!NOTE]")
            for ln in el["lines"]:
                lines.append(f"> {ln}")
        elif t == "image":
            lines.append(f"![]({el['src']})")
    p.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")
    return p


def write_txt(topic_dir: Path, title: str, elements: list[dict]) -> Path:
    topic_dir.mkdir(parents=True, exist_ok=True)
    p = topic_dir / "tema.txt"
    lines: list[str] = [title, ""]
    for el in elements:
        t = el["type"]
        if t == "page_marker":
            lines.append("")
            lines.append(f"[Página {el['page']}]\n")
        elif t == "heading":
            lines.append(el.get("text", ""))
        elif t == "paragraph":
            lines.append(el.get("text", ""))
        elif t == "list":
            for it in el.get("items", []):
                lines.append(f"- {it}")
        elif t == "callout":
            lines.append("[RECUADRO]")
            for ln in el.get("lines", []):
                lines.append(ln)
        elif t == "image":
            lines.append(f"[IMAGEN] {el.get('src')}")
    p.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")
    return p


def write_json(topic_dir: Path, title: str, meta: dict, elements: list[dict]) -> Path:
    topic_dir.mkdir(parents=True, exist_ok=True)
    p = topic_dir / "tema.json"
    payload = {
        "title": title,
        "meta": meta,
        "elements": elements,
    }
    p.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return p


def _html_escape(s: str) -> str:
    return (
        s.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def clean_xml_text(s: str) -> str:
    # Remove ASCII control chars that are invalid in XML 1.0
    # Allowed: \t, \n, \r. Disallowed: 0x00-0x08, 0x0B, 0x0C, 0x0E-0x1F
    return re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", s)


def write_html(topic_dir: Path, title: str, elements: list[dict]) -> Path:
    topic_dir.mkdir(parents=True, exist_ok=True)
    p = topic_dir / "tema.html"
    css = """
    body{font-family:system-ui,-apple-system,Segoe UI,Roboto,Ubuntu,Cantarell,Noto Sans,sans-serif;line-height:1.5;max-width:980px;margin:24px auto;padding:0 16px;}
    h1,h2,h3{line-height:1.2}
    img{max-width:100%;height:auto}
    .callout{border-left:4px solid #444;background:#f6f6f6;padding:12px 14px;margin:12px 0}
    .muted{color:#666;font-size:0.9em}
    """.strip()

    html_parts = [
        "<!doctype html>",
        "<html lang='es'>",
        "<head>",
        "<meta charset='utf-8'>",
        "<meta name='viewport' content='width=device-width,initial-scale=1'>",
        f"<title>{_html_escape(title)}</title>",
        f"<style>{css}</style>",
        "</head>",
        "<body>",
        f"<h1>{_html_escape(title)}</h1>",
    ]

    in_ul = False

    def close_ul():
        nonlocal in_ul
        if in_ul:
            html_parts.append("</ul>")
            in_ul = False

    for el in elements:
        t = el["type"]
        if t == "page_marker":
            close_ul()
            html_parts.append(f"<p class='muted'>Página {int(el['page'])}</p>")
        elif t == "heading":
            close_ul()
            level = int(el.get("level", 2))
            level = max(2, min(6, level))
            html_parts.append(f"<h{level}>{_html_escape(el.get('text',''))}</h{level}>")
        elif t == "paragraph":
            close_ul()
            html_parts.append(f"<p>{_html_escape(el.get('text',''))}</p>")
        elif t == "list":
            if not in_ul:
                html_parts.append("<ul>")
                in_ul = True
            for it in el.get("items", []):
                html_parts.append(f"<li>{_html_escape(it)}</li>")
        elif t == "callout":
            close_ul()
            lines = el.get("lines", [])
            html_parts.append("<div class='callout'>")
            for ln in lines:
                html_parts.append(f"<p>{_html_escape(ln)}</p>")
            html_parts.append("</div>")
        elif t == "image":
            close_ul()
            src = el.get("src", "")
            html_parts.append(f"<img src='{_html_escape(src)}' alt=''>")

    close_ul()
    html_parts.append("</body></html>")
    p.write_text("\n".join(html_parts) + "\n", encoding="utf-8")
    return p


def write_docx(topic_dir: Path, title: str, elements: list[dict]) -> Path:
    if Document is None:
        raise RuntimeError("Falta dependencia: python-docx. Instale con `pip install python-docx`.")
    topic_dir.mkdir(parents=True, exist_ok=True)
    p = topic_dir / "tema.docx"
    doc = Document()
    doc.add_heading(clean_xml_text(title), level=1)

    for el in elements:
        t = el["type"]
        if t == "page_marker":
            doc.add_paragraph(clean_xml_text(f"Página {int(el['page'])}")).italic = True
        elif t == "heading":
            level = int(el.get("level", 2))
            level = max(2, min(6, level))
            doc.add_heading(clean_xml_text(el.get("text", "")), level=level)
        elif t == "paragraph":
            doc.add_paragraph(clean_xml_text(el.get("text", "")))
        elif t == "list":
            for it in el.get("items", []):
                doc.add_paragraph(clean_xml_text(it), style="List Bullet")
        elif t == "callout":
            for ln in el.get("lines", []):
                doc.add_paragraph(clean_xml_text(ln), style="Intense Quote")
        elif t == "image":
            src = el.get("src")
            if not src:
                continue
            img_path = topic_dir / src
            if img_path.exists():
                # Ancho aproximado; Word ajusta altura automáticamente
                try:
                    doc.add_picture(str(img_path), width=Inches(5.8))
                except UnrecognizedImageError:
                    # Algunas imágenes extraídas pueden no ser reconocidas por docx.
                    # Intentamos convertir a PNG con Pillow si está disponible.
                    if Image is None:
                        continue
                    try:
                        converted_dir = (topic_dir / "images" / "_converted")
                        converted_dir.mkdir(parents=True, exist_ok=True)
                        converted_path = converted_dir / (img_path.stem + ".png")
                        with Image.open(img_path) as im:
                            im.save(converted_path, format="PNG")
                        doc.add_picture(str(converted_path), width=Inches(5.8))
                    except Exception:
                        continue

    doc.save(str(p))
    return p


def write_pdf(topic_dir: Path, title: str, src_doc: fitz.Document, start_page_idx: int, end_page_idx: int) -> Path:
    topic_dir.mkdir(parents=True, exist_ok=True)
    p = topic_dir / "tema.pdf"
    out = fitz.open()
    out.insert_pdf(src_doc, from_page=start_page_idx, to_page=end_page_idx)
    out.save(str(p))
    out.close()
    return p


def detect_heading_topics(
    doc: fitz.Document,
    start_i: int,
    end_i: int,
    topic_re: re.Pattern,
    heading_scale: float,
) -> list[tuple[str, int]]:
    found: list[tuple[str, int]] = []
    seen: set[str] = set()
    for page_i in range(start_i, end_i + 1):
        items = collect_items(doc, page_i)
        sizes = [it.font_size for it in items if it.kind == "text" and it.font_size]
        base = _median([float(s) for s in sizes if s]) if sizes else 0.0
        for it in items:
            if it.kind != "text" or not it.text:
                continue
            raw = normalize_title(it.text)
            if is_probably_footer_or_noise(raw):
                continue
            lvl = heading_level(float(it.font_size or 0.0), base, heading_scale)
            if not lvl:
                continue
            if not topic_re.match(raw):
                continue
            key = raw.lower()
            if key in seen:
                continue
            seen.add(key)
            found.append((raw, page_i))
    found.sort(key=lambda t: t[1])
    return found


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("pdf", help="Ruta al PDF")
    ap.add_argument("--outdir", default="output_all", help="Directorio de salida")
    ap.add_argument(
        "--formats",
        default="md",
        help="Formatos: md,html,docx,pdf,txt,json o 'all' (separados por coma)",
    )
    ap.add_argument(
        "--mode",
        default="auto",
        choices=["auto", "headings", "toc"],
        help="Cómo segmentar temas: auto (headings->toc), headings, toc",
    )
    ap.add_argument(
        "--topic-regex",
        default=r"^(tema|unidad|cap[ií]tulo)\s*\d+\b",
        help="Regex (case-insensitive) para detectar inicio de tema por encabezado.",
    )
    ap.add_argument("--heading-scale", type=float, default=1.6)
    ap.add_argument("--start-page", type=int, default=1)
    ap.add_argument("--end-page", type=int, default=None)
    ap.add_argument("--toc-max-pages", type=int, default=40)
    args = ap.parse_args()

    pdf_path = Path(args.pdf)
    if not pdf_path.exists():
        raise SystemExit(f"Archivo no encontrado: {pdf_path}")

    formats_raw = [f.strip().lower() for f in args.formats.split(",") if f.strip()]
    if "all" in formats_raw:
        formats = {"md", "html", "docx", "pdf", "txt", "json"}
    else:
        formats = set(formats_raw)

    out_root = Path(args.outdir) / pdf_path.stem
    out_root.mkdir(parents=True, exist_ok=True)

    doc = fitz.open(str(pdf_path))
    start_i = max(0, args.start_page - 1)
    end_i = (args.end_page - 1) if args.end_page else (doc.page_count - 1)
    end_i = min(end_i, doc.page_count - 1)

    topic_re = re.compile(args.topic_regex, re.IGNORECASE)

    topics: list[tuple[str, int]] = []

    if args.mode in ("auto", "headings"):
        topics = detect_heading_topics(doc, start_i, end_i, topic_re, args.heading_scale)

    if (not topics) and args.mode in ("auto", "toc"):
        toc = parse_toc_entries(doc, max_pages=args.toc_max_pages)
        offset = estimate_printed_to_pdf_offset(doc)
        if offset is not None:
            for title, printed in toc:
                pdf_page = printed + offset
                if 1 <= pdf_page <= doc.page_count:
                    topics.append((title, pdf_page - 1))
        topics.sort(key=lambda t: t[1])

    if not topics:
        topics = [(pdf_path.stem, start_i)]

    # Make ranges
    ranges: list[dict] = []
    for idx, (title, start_page_idx) in enumerate(topics, start=1):
        end_page_idx = (topics[idx][1] - 1) if idx < len(topics) else end_i
        start_page_idx = max(start_i, start_page_idx)
        end_page_idx = min(end_i, end_page_idx)
        if end_page_idx < start_page_idx:
            continue
        ranges.append(
            {
                "index": idx,
                "title": normalize_title(title),
                "slug": slugify(title),
                "start_page": start_page_idx + 1,
                "end_page": end_page_idx + 1,
                "start_page_idx": start_page_idx,
                "end_page_idx": end_page_idx,
            }
        )

    index_md_lines = ["# Índice de temas", ""]
    index_json: list[dict] = []

    for t in ranges:
        idx = int(t["index"])
        title = t["title"]
        slug = t["slug"]
        topic_dir = out_root / f"tema_{idx:03d}_{slug}"
        topic_dir.mkdir(parents=True, exist_ok=True)

        meta = {
            "source_pdf": pdf_path.name,
            "start_page": t["start_page"],
            "end_page": t["end_page"],
        }

        elements: list[dict] = []
        images_meta: list[dict] = []
        if formats.intersection({"md", "html", "docx", "txt", "json"}):
            elements, images_meta = extract_topic_assets(
                doc,
                range(t["start_page_idx"], t["end_page_idx"] + 1),
                topic_dir,
                args.heading_scale,
            )
            meta["images"] = images_meta

        produced: dict[str, str] = {}
        if "md" in formats:
            produced["md"] = str(write_md(topic_dir, title, elements).relative_to(out_root))
        if "html" in formats:
            produced["html"] = str(write_html(topic_dir, title, elements).relative_to(out_root))
        if "txt" in formats:
            produced["txt"] = str(write_txt(topic_dir, title, elements).relative_to(out_root))
        if "json" in formats:
            produced["json"] = str(write_json(topic_dir, title, meta, elements).relative_to(out_root))
        if "docx" in formats:
            produced["docx"] = str(write_docx(topic_dir, title, elements).relative_to(out_root))
        if "pdf" in formats:
            produced["pdf"] = str(
                write_pdf(topic_dir, title, doc, t["start_page_idx"], t["end_page_idx"]).relative_to(out_root)
            )

        # Index entry
        link_target = produced.get("md") or produced.get("html") or produced.get("txt") or produced.get("pdf")
        if link_target:
            index_md_lines.append(f"- [{title}]({link_target.replace(os.sep, '/')})")
        else:
            index_md_lines.append(f"- {title}")

        index_json.append({"title": title, "meta": meta, "files": produced})

    (out_root / "index.md").write_text("\n".join(index_md_lines) + "\n", encoding="utf-8")
    (out_root / "index.json").write_text(json.dumps(index_json, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    doc.close()
    print(f"Listo. Exportado en: {out_root}")


if __name__ == "__main__":
    main()
